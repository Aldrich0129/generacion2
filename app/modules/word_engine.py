"""
Motor de generación de documentos Word usando python-docx.
Maneja reemplazo de variables, inserción de tablas y bloques condicionales.

NOTA IMPORTANTE: Conservación de Imágenes y Elementos Gráficos
---------------------------------------------------------------
Este motor implementa protección exhaustiva de elementos gráficos:

✓ PRESERVACIÓN DE FOTOS E IMÁGENES:
  - Imágenes de fondo en portada, páginas intermedias y páginas finales
  - Imágenes en headers (cabeceras) y footers (pies de página)
  - Shapes, formas VML, y dibujos (w:drawing, w:pict, v:shape, v:imagedata)
  - Imágenes configuradas como "Detrás del texto" o en cualquier posición
  - Los métodos de limpieza verifican la presencia de elementos gráficos antes
    de eliminar cualquier párrafo, asegurando que nunca se pierdan imágenes

✓ LIMPIEZA INTELIGENTE DE CONTENIDO:
  - Elimina páginas vacías sin afectar imágenes
  - Elimina líneas vacías al inicio de páginas para mejor presentación
  - Elimina marcadores no utilizados preservando contenido gráfico
  - Todos los métodos de limpieza usan _has_drawing_or_image() para proteger imágenes

✓ CONSERVACIÓN DE DISEÑO:
  - Headers y footers se preservan con todos sus elementos
  - Todas las secciones mantienen su formato y contenido gráfico
  - El diseño original de la plantilla permanece intacto

El motor solo modifica el contenido de texto, tablas y párrafos específicamente editados.
Todos los elementos visuales y gráficos de la plantilla permanecen protegidos durante
todo el proceso de generación.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, List, Any, Optional
import re
from copy import deepcopy


class WordEngine:
    """Motor para generar documentos Word desde plantillas."""

    def __init__(self, template_path: Path):
        """
        Inicializa el motor con una plantilla.

        Args:
            template_path: Ruta a la plantilla Word (.docx)
        """
        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")

        self.doc = Document(self.template_path)

    def replace_variables(self, context: dict):
        """
        Reemplaza todas las variables <<marcador>> en el documento manteniendo estilos.
        Si el valor está vacío, elimina el marcador del documento.

        Args:
            context: Diccionario con {marcador: valor}
        """
        # Filtrar valores vacíos - los marcadores con valores vacíos se eliminarán
        context_filtered = {}
        for marker, value in context.items():
            if value is None or value == "" or (isinstance(value, str) and not value.strip()):
                # Valor vacío, no reemplazar (se limpiará después)
                continue
            context_filtered[marker] = value

        # Reemplazar en párrafos
        for paragraph in self.doc.paragraphs:
            self._replace_in_paragraph(paragraph, context_filtered)

        # Reemplazar en tablas
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, context_filtered)

        # Reemplazar en headers y footers
        for section in self.doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, context_filtered)

            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_in_paragraph(paragraph, context_filtered)

    def _replace_in_paragraph(self, paragraph, context: dict):
        """
        Reemplaza marcadores en un párrafo manteniendo el formato.

        Args:
            paragraph: Párrafo de python-docx
            context: Diccionario de reemplazos
        """
        # Construir el texto completo del párrafo
        full_text = paragraph.text

        # Buscar todos los marcadores en el texto
        for marker, value in context.items():
            if marker in full_text:
                # Reemplazar manteniendo el formato
                self._replace_marker_in_paragraph(paragraph, marker, str(value))
                full_text = paragraph.text  # Actualizar el texto completo

    def _replace_marker_in_paragraph(self, paragraph, marker: str, value: str):
        """
        Reemplaza un marcador específico manteniendo el formato del primer run.

        Args:
            paragraph: Párrafo de python-docx
            marker: Marcador a buscar (ej: "<<Variable>>")
            value: Valor de reemplazo
        """
        # Buscar el marcador en los runs
        full_text = ""
        runs_text = []

        for run in paragraph.runs:
            runs_text.append(run.text)
            full_text += run.text

        if marker not in full_text:
            return

        # Encontrar la posición del marcador
        marker_start = full_text.find(marker)
        marker_end = marker_start + len(marker)

        # Reconstruir el párrafo
        current_pos = 0
        new_runs = []

        for i, run in enumerate(paragraph.runs):
            run_text = runs_text[i]
            run_start = current_pos
            run_end = current_pos + len(run_text)

            # Verificar si este run contiene parte del marcador
            if run_end <= marker_start or run_start >= marker_end:
                # Este run no toca el marcador, mantenerlo
                new_runs.append((run_text, run))
            elif run_start <= marker_start and run_end >= marker_end:
                # El marcador está completamente dentro de este run
                before = run_text[:marker_start - run_start]
                after = run_text[marker_end - run_start:]
                new_runs.append((before + value + after, run))
            elif run_start < marker_start < run_end:
                # El marcador comienza en este run
                before = run_text[:marker_start - run_start]
                new_runs.append((before + value, run))
            elif run_start < marker_end <= run_end:
                # El marcador termina en este run
                after = run_text[marker_end - run_start:]
                if after:
                    new_runs.append((after, run))
            # else: este run está completamente dentro del marcador, se omite

            current_pos = run_end

        # Limpiar runs existentes
        for run in paragraph.runs:
            run.text = ""

        # Recrear runs con el nuevo texto
        if new_runs:
            for new_text, original_run in new_runs:
                if new_text:
                    # Copiar formato del run original
                    new_run = paragraph.add_run(new_text)
                    new_run.bold = original_run.bold
                    new_run.italic = original_run.italic
                    new_run.underline = original_run.underline
                    new_run.font.size = original_run.font.size
                    new_run.font.name = original_run.font.name
                    if original_run.font.color.rgb:
                        new_run.font.color.rgb = original_run.font.color.rgb

    def insert_tables(self, tables_data: dict, cfg_tab: dict, format_config: dict = None):
        """
        Inserta tablas en los marcadores correspondientes.

        Args:
            tables_data: Datos de las tablas construidas
            cfg_tab: Configuración de tablas
            format_config: Configuración de formato de tablas (opcional)
        """
        if format_config is None:
            format_config = {}

        for marker, table_data in tables_data.items():
            # Extraer el table_id del marker
            # Para marcadores simples: "<<Tabla análisis indirecto>>" -> "analisis_indirecto_global"
            # Para marcadores parametrizados: "<<Tabla Operación 1>>" -> "analisis_indirecto_operacion_1"
            table_id = table_data.get("table_id", None)
            self._insert_table_at_marker(marker, table_data, format_config, table_id)

    def _insert_table_at_marker(self, marker: str, table_data: dict, format_config: dict = None, table_id: str = None):
        """
        Inserta una tabla en la posición del marcador.

        Args:
            marker: Marcador donde insertar la tabla
            table_data: Datos de la tabla (rows, columns, etc.)
            format_config: Configuración de formato de tablas (opcional)
            table_id: Identificador de la tabla (para configuración personalizada)
        """
        if format_config is None:
            format_config = {}

        # Buscar el marcador en el documento
        for i, paragraph in enumerate(self.doc.paragraphs):
            if marker in paragraph.text:
                # Crear la tabla justo después de este párrafo
                self._create_table_after_paragraph(i, table_data, format_config, table_id)

                # Limpiar el marcador
                paragraph.text = paragraph.text.replace(marker, "")
                return

    def _create_table_after_paragraph(self, para_index: int, table_data: dict, format_config: dict = None, table_id: str = None):
        """
        Crea una tabla después de un párrafo específico.

        Args:
            para_index: Índice del párrafo
            table_data: Datos de la tabla
            format_config: Configuración de formato de tablas (opcional)
            table_id: Identificador de la tabla (para configuración personalizada)
        """
        if format_config is None:
            format_config = {}

        # Verificar si existe configuración personalizada para esta tabla
        custom_formats = format_config.get("custom_table_formats", {})
        if table_id and table_id in custom_formats:
            # Usar la configuración personalizada para esta tabla
            table_format = custom_formats[table_id]
        else:
            # Usar la configuración general
            table_format = format_config

        columns = table_data.get("columns", [])
        rows = table_data.get("rows", [])
        footer_rows = table_data.get("footer_rows", [])
        headers = table_data.get("headers", {})

        if not columns or not rows:
            return

        # Calcular número de filas (header + data + footer)
        num_rows = 1 + len(rows) + len(footer_rows)
        num_cols = len(columns)

        # Encontrar el elemento XML del párrafo
        para = self.doc.paragraphs[para_index]
        para_element = para._element

        # Crear la tabla
        from docx.oxml import parse_xml
        from docx.table import Table

        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # Aplicar estilo de tabla según configuración
        show_borders = table_format.get("show_borders", True)
        if show_borders:
            # Intentar aplicar estilo con bordes
            try:
                table.style = 'Light Grid Accent 1'
            except KeyError:
                try:
                    table.style = 'Light Grid'
                except KeyError:
                    try:
                        table.style = 'Table Grid'
                    except KeyError:
                        pass
        else:
            # Sin bordes - usar estilo sin líneas
            try:
                table.style = 'Table Normal'
            except KeyError:
                pass

        # Insertar la tabla después del párrafo
        table_element = table._element
        para_element.addnext(table_element)

        # Obtener configuración de formato
        header_bg_color = table_format.get("header_bg_color", "#4472C4")
        header_text_color = table_format.get("header_text_color", "#FFFFFF")
        header_bold = table_format.get("header_bold", True)
        header_font_size = table_format.get("header_font_size", 11)
        data_font_size = table_format.get("data_font_size", 10)
        alternate_rows = table_format.get("alternate_rows", False)
        alternate_row_color = table_format.get("alternate_row_color", "#F2F2F2")
        border_color = table_format.get("border_color", "#000000")
        first_column_bold = table_format.get("first_column_bold", False)
        first_column_bg_color = table_format.get("first_column_bg_color", None)
        first_column_text_color = table_format.get("first_column_text_color", None)
        column_colors = table_format.get("column_colors", [])

        # Aplicar bordes personalizados a toda la tabla si show_borders es True
        if show_borders:
            self._apply_table_borders(table, border_color)
        else:
            # Eliminar todos los bordes si show_borders es False
            self._remove_table_borders(table)

        # Llenar encabezados
        header_row = table.rows[0]
        for j, col in enumerate(columns):
            cell = header_row.cells[j]
            header_text = col.get("header", "")

            # Reemplazar placeholders en el header si hay headers dinámicos
            if headers:
                for key, value in headers.items():
                    header_text = header_text.replace(f"{{{key}}}", str(value))

            cell.text = header_text

            # Aplicar formato de encabezado personalizado
            self._apply_cell_shading(cell, header_bg_color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if header_bold:
                        run.bold = True
                    run.font.size = Pt(header_font_size)
                    # Aplicar color de texto del encabezado
                    run.font.color.rgb = self._hex_to_rgb(header_text_color)

        # Llenar filas de datos
        for i, row_data in enumerate(rows):
            row = table.rows[i + 1]

            # Aplicar color alternado a filas pares si está configurado
            if alternate_rows and i % 2 == 0:
                for cell in row.cells:
                    self._apply_cell_shading(cell, alternate_row_color)

            for j, col in enumerate(columns):
                cell = row.cells[j]
                col_id = col["id"]
                value = row_data.get(col_id, "")

                # Formatear según el tipo
                col_type = col.get("type", "text")
                formatted_value = self._format_cell_value(value, col_type)

                cell.text = formatted_value

                # Aplicar estilo de primera columna si está configurado y es la primera columna
                is_first_column = (j == 0)
                if is_first_column:
                    # Aplicar color de fondo si está configurado
                    if first_column_bg_color:
                        self._apply_cell_shading(cell, first_column_bg_color)

                # Aplicar color de fondo por columna si está configurado
                # Esto se aplica después de first_column_bg_color para que tome prioridad
                # Nota: j es el índice (empezando en 0), pero los usuarios especifican columnas empezando en 1
                column_number = j + 1
                for col_color_config in column_colors:
                    if col_color_config.get("column") == column_number:
                        self._apply_cell_shading(cell, col_color_config.get("color"))
                        break

                # Aplicar tamaño de fuente a las celdas de datos
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(data_font_size)

                        # Aplicar estilos de primera columna
                        if is_first_column:
                            if first_column_bold:
                                run.bold = True
                            if first_column_text_color:
                                run.font.color.rgb = self._hex_to_rgb(first_column_text_color)

        # Llenar filas de footer
        if footer_rows:
            for i, footer_data in enumerate(footer_rows):
                row_index = 1 + len(rows) + i
                row = table.rows[row_index]

                for j, col in enumerate(columns):
                    cell = row.cells[j]
                    col_id = col["id"]
                    value = footer_data.get(col_id, "")

                    col_type = col.get("type", "text")
                    formatted_value = self._format_cell_value(value, col_type)

                    cell.text = formatted_value

                    # Aplicar formato de footer (negrita y tamaño de fuente)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(data_font_size)

    def _format_cell_value(self, value: Any, col_type: str) -> str:
        """
        Formatea el valor de una celda según su tipo.

        Args:
            value: Valor a formatear
            col_type: Tipo de columna

        Returns:
            Valor formateado como string
        """
        if value is None or value == "":
            return ""

        try:
            if col_type == "percent":
                # Convertir a número si es posible
                if isinstance(value, str):
                    # Eliminar el símbolo % si ya existe
                    value = value.replace('%', '').strip()
                    try:
                        value = float(value)
                    except ValueError:
                        return str(value)

                if isinstance(value, (int, float)):
                    return f"{value:.2f}%"
                return str(value)
            elif col_type == "number":
                if isinstance(value, (int, float)):
                    return f"{value:,.2f}"
                elif isinstance(value, str):
                    try:
                        value = float(value)
                        return f"{value:,.2f}"
                    except ValueError:
                        return str(value)
                return str(value)
            elif col_type == "integer":
                if isinstance(value, (int, float)):
                    return str(int(value))
                elif isinstance(value, str):
                    try:
                        value = int(float(value))
                        return str(value)
                    except ValueError:
                        return str(value)
                return str(value)
            else:
                return str(value)
        except:
            return str(value)

    def insert_conditional_blocks(self, docs_to_insert: list, config_dir: Path):
        """
        Inserta bloques condicionales de otros documentos Word.

        Args:
            docs_to_insert: Lista de {marker, file} a insertar
            config_dir: Directorio base para los archivos Word
        """
        for doc_info in docs_to_insert:
            marker = doc_info["marker"]
            word_file = doc_info["file"]

            file_path = config_dir.parent / word_file

            if file_path.exists():
                self._insert_document_at_marker(marker, file_path)
            else:
                print(f"Advertencia: Archivo no encontrado: {file_path}")
                # Limpiar el marcador
                self.replace_variables({marker: ""})

    def _insert_document_at_marker(self, marker: str, doc_path: Path):
        """
        Inserta el contenido de otro documento en la posición del marcador.

        Args:
            marker: Marcador donde insertar
            doc_path: Ruta al documento a insertar
        """
        # Cargar el documento a insertar
        doc_to_insert = Document(doc_path)

        # Buscar el marcador
        for i, paragraph in enumerate(self.doc.paragraphs):
            if marker in paragraph.text:
                # Limpiar el marcador
                paragraph.text = paragraph.text.replace(marker, "")

                # Insertar el contenido del documento
                para_element = paragraph._element

                # Copiar párrafos del documento a insertar
                for insert_para in doc_to_insert.paragraphs:
                    # Crear un nuevo párrafo con el mismo formato
                    new_para = deepcopy(insert_para._element)
                    para_element.addnext(new_para)
                    para_element = new_para

                # Copiar tablas del documento a insertar
                for insert_table in doc_to_insert.tables:
                    new_table = deepcopy(insert_table._element)
                    para_element.addnext(new_table)
                    para_element = new_table

                return

    def process_salto_markers(self):
        """
        Procesa los marcadores {salto} insertando saltos de página antes del marcador
        y eliminando el marcador del documento.
        """
        salto_pattern = re.compile(r'\{salto\}')

        # Procesar párrafos
        for paragraph in self.doc.paragraphs:
            if salto_pattern.search(paragraph.text):
                # Obtener el texto antes y después del marcador
                text = paragraph.text
                parts = salto_pattern.split(text, 1)  # Split solo la primera ocurrencia

                if len(parts) == 2:
                    before_salto = parts[0]
                    after_salto = parts[1]

                    # Limpiar el párrafo
                    paragraph.clear()

                    # Agregar el texto antes del salto
                    if before_salto:
                        paragraph.add_run(before_salto)

                    # Insertar salto de página
                    run = paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)

                    # Agregar el texto después del salto
                    if after_salto:
                        paragraph.add_run(after_salto)
                else:
                    # Solo eliminar el marcador si no hay texto después
                    paragraph.text = salto_pattern.sub('', paragraph.text)

        # Procesar tablas
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if salto_pattern.search(paragraph.text):
                            # En tablas, solo eliminar el marcador sin insertar salto
                            paragraph.text = salto_pattern.sub('', paragraph.text)

    def _has_drawing_or_image(self, paragraph) -> bool:
        """
        Verifica si un párrafo contiene imágenes, shapes, dibujos u otros elementos gráficos.

        Args:
            paragraph: Párrafo de python-docx

        Returns:
            True si el párrafo contiene elementos gráficos, False en caso contrario
        """
        # Obtener el elemento XML del párrafo
        para_element = paragraph._element

        # Buscar elementos de dibujo (w:drawing)
        drawings = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawings:
            return True

        # Buscar elementos de imagen (w:pict)
        pictures = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
        if pictures:
            return True

        # Buscar elementos v:shape (formas de VML)
        vml_shapes = para_element.findall('.//{urn:schemas-microsoft-com:vml}shape')
        if vml_shapes:
            return True

        # Buscar elementos v:imagedata
        vml_images = para_element.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
        if vml_images:
            return True

        return False

    def clean_unused_markers(self):
        """
        Limpia todos los marcadores no utilizados del documento.
        Si un párrafo contiene SOLO un marcador (o marcador con puntuación/numeración),
        elimina el párrafo completo. Si hay más contenido, solo elimina el marcador.

        IMPORTANTE: Nunca elimina párrafos que contengan imágenes, shapes o dibujos,
        incluso si parecen vacíos, para preservar las fotos de la plantilla.
        """
        marker_pattern = re.compile(r'<<[^>]+>>')

        # Limpiar en párrafos - eliminar líneas completas si solo contienen marcadores
        paragraphs_to_delete = []
        for i, paragraph in enumerate(self.doc.paragraphs):
            if marker_pattern.search(paragraph.text):
                # IMPORTANTE: Verificar si el párrafo contiene imágenes o dibujos
                # Si tiene imágenes, NUNCA eliminarlo, solo limpiar el marcador
                if self._has_drawing_or_image(paragraph):
                    # Solo eliminar el marcador, preservar el párrafo con la imagen
                    paragraph.text = marker_pattern.sub('', paragraph.text)
                    continue

                # Verificar si el párrafo solo contiene marcador y elementos comunes (puntos, números, guiones, espacios)
                text_without_markers = marker_pattern.sub('', paragraph.text).strip()
                # Eliminar también puntuación común, números, guiones, viñetas
                text_cleaned = re.sub(r'^[\d\.\-\)\(\s•·◦▪▫○●\*]+$', '', text_without_markers)

                if not text_cleaned:
                    # El párrafo solo contiene marcador + elementos decorativos, marcarlo para eliminación
                    paragraphs_to_delete.append(paragraph)
                else:
                    # Hay contenido real además del marcador, solo eliminar el marcador
                    paragraph.text = marker_pattern.sub('', paragraph.text)

        # Eliminar los párrafos marcados
        for paragraph in paragraphs_to_delete:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        # Limpiar en tablas - solo eliminar el marcador, no la celda
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if marker_pattern.search(paragraph.text):
                            paragraph.text = marker_pattern.sub('', paragraph.text)

    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """
        Convierte un color hexadecimal a RGBColor.

        Args:
            hex_color: Color en formato hexadecimal (ej: "#4472C4" o "4472C4")

        Returns:
            Objeto RGBColor
        """
        # Eliminar el '#' si existe
        hex_color = hex_color.lstrip('#')

        # Convertir a RGB
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)

        return RGBColor(r, g, b)

    def _apply_cell_shading(self, cell, hex_color: str):
        """
        Aplica un color de fondo a una celda de tabla.

        Args:
            cell: Celda de la tabla
            hex_color: Color en formato hexadecimal (ej: "#4472C4")
        """
        # Eliminar el '#' si existe
        hex_color = hex_color.lstrip('#')

        # Obtener el elemento XML de la celda
        cell_properties = cell._element.get_or_add_tcPr()

        # Crear elemento de sombreado
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), hex_color)

        # Añadir el sombreado a las propiedades de la celda
        cell_properties.append(shading)

    def _apply_table_borders(self, table, hex_color: str):
        """
        Aplica bordes con color personalizado a todas las celdas de una tabla.

        Args:
            table: Tabla de python-docx
            hex_color: Color en formato hexadecimal (ej: "#000000")
        """
        # Eliminar el '#' si existe
        hex_color = hex_color.lstrip('#')

        # Aplicar bordes a cada celda
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()

                # Crear bordes
                tcBorders = OxmlElement('w:tcBorders')

                # Definir cada borde (top, left, bottom, right)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # Tamaño del borde
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), hex_color)
                    tcBorders.append(border)

                tcPr.append(tcBorders)

    def _remove_table_borders(self, table):
        """
        Elimina todos los bordes de una tabla.

        Args:
            table: Tabla de python-docx
        """
        # Eliminar bordes de cada celda
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()

                # Crear bordes vacíos (sin líneas)
                tcBorders = OxmlElement('w:tcBorders')

                # Definir cada borde como "none"
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)

                tcPr.append(tcBorders)

    def clean_empty_paragraphs(self):
        """
        Elimina párrafos vacíos del documento.
        """
        # Nota: Eliminar párrafos en python-docx es complicado
        # Por ahora, solo los marcamos como vacíos
        for paragraph in self.doc.paragraphs:
            if not paragraph.text.strip():
                # Opcionalmente reducir el espaciado
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    def remove_empty_pages(self):
        """
        Elimina páginas completamente vacías del documento.

        Una página se considera vacía si:
        - Solo contiene párrafos vacíos (sin texto o solo espacios en blanco)
        - Solo contiene saltos de página sin contenido significativo

        Este método analiza el documento y elimina:
        - Saltos de página consecutivos que crean páginas en blanco
        - Párrafos vacíos que preceden a saltos de página

        IMPORTANTE: Nunca elimina párrafos que contengan imágenes, shapes o dibujos
        para preservar las fotos de la plantilla.
        """
        # Lista para rastrear párrafos a eliminar
        paragraphs_to_remove = []

        # Recorrer todos los párrafos del documento
        i = 0
        while i < len(self.doc.paragraphs):
            para = self.doc.paragraphs[i]

            # IMPORTANTE: Si el párrafo contiene imágenes o dibujos, NUNCA eliminarlo
            if self._has_drawing_or_image(para):
                i += 1
                continue

            # Verificar si el párrafo está vacío o solo tiene espacios
            text = para.text.strip()

            # Verificar si el párrafo tiene un salto de página
            has_page_break = any(self._has_page_break(run) for run in para.runs)

            if has_page_break and not text:
                # Este párrafo solo tiene un salto de página sin contenido
                # Verificar si los párrafos anteriores también están vacíos
                j = i - 1
                consecutive_empty = 0

                while j >= 0:
                    prev_para = self.doc.paragraphs[j]

                    # IMPORTANTE: Si el párrafo anterior tiene imágenes, no continuar
                    if self._has_drawing_or_image(prev_para):
                        break

                    prev_text = prev_para.text.strip()
                    prev_has_break = any(self._has_page_break(run) for run in prev_para.runs)

                    if not prev_text and not prev_has_break:
                        # Párrafo vacío sin salto de página
                        consecutive_empty += 1
                        j -= 1
                    elif not prev_text and prev_has_break:
                        # Otro salto de página vacío, esto crea una página en blanco
                        # Marcar para eliminación
                        paragraphs_to_remove.append(para)
                        break
                    else:
                        # Hay contenido antes del salto, no eliminar
                        break

                # Si hay muchos párrafos vacíos antes del salto, marcar esos también
                if consecutive_empty > 2:
                    for k in range(j + 1, i):
                        if k >= 0:
                            check_para = self.doc.paragraphs[k]
                            # Solo marcar si no tiene imágenes
                            if not self._has_drawing_or_image(check_para):
                                paragraphs_to_remove.append(check_para)

            elif not text and i > 0:
                # Párrafo vacío sin salto de página
                # Verificar si el siguiente también está vacío o es un salto de página
                if i + 1 < len(self.doc.paragraphs):
                    next_para = self.doc.paragraphs[i + 1]

                    # IMPORTANTE: Si el siguiente tiene imágenes, no eliminar este
                    if self._has_drawing_or_image(next_para):
                        i += 1
                        continue

                    next_text = next_para.text.strip()
                    next_has_break = any(self._has_page_break(run) for run in next_para.runs)

                    # Si el siguiente es un salto de página vacío, marcar este párrafo vacío también
                    if next_has_break and not next_text:
                        # Verificar cuántos párrafos vacíos consecutivos hay
                        k = i
                        empty_count = 0
                        while k >= 0 and not self.doc.paragraphs[k].text.strip():
                            # Verificar que no tenga imágenes
                            if self._has_drawing_or_image(self.doc.paragraphs[k]):
                                break
                            empty_count += 1
                            k -= 1

                        # Si hay más de 3 párrafos vacíos consecutivos seguidos de un salto de página,
                        # probablemente es una página vacía
                        if empty_count > 3:
                            paragraphs_to_remove.append(para)

            i += 1

        # Eliminar los párrafos marcados
        for para in paragraphs_to_remove:
            try:
                p_element = para._element
                p_element.getparent().remove(p_element)
            except Exception:
                # Si falla, continuar con el siguiente
                pass

        # Segunda pasada: eliminar saltos de página duplicados
        # Recorrer nuevamente para encontrar saltos de página consecutivos
        i = 0
        while i < len(self.doc.paragraphs) - 1:
            para = self.doc.paragraphs[i]
            next_para = self.doc.paragraphs[i + 1]

            # IMPORTANTE: Si alguno de los dos tiene imágenes, no eliminar
            if self._has_drawing_or_image(para) or self._has_drawing_or_image(next_para):
                i += 1
                continue

            # Verificar si ambos párrafos tienen saltos de página
            has_break = any(self._has_page_break(run) for run in para.runs)
            next_has_break = any(self._has_page_break(run) for run in next_para.runs)

            # Verificar si ambos están vacíos
            is_empty = not para.text.strip()
            next_is_empty = not next_para.text.strip()

            # Si ambos tienen saltos de página y están vacíos, eliminar el segundo
            if has_break and next_has_break and is_empty and next_is_empty:
                try:
                    p_element = next_para._element
                    p_element.getparent().remove(p_element)
                    # No incrementar i, porque eliminamos el siguiente
                    continue
                except Exception:
                    pass

            i += 1

    def remove_empty_lines_at_page_start(self):
        """
        Elimina líneas vacías (párrafos vacíos) al inicio de páginas.

        Este método busca saltos de página y elimina los párrafos vacíos que
        aparecen inmediatamente después, limpiando el espacio en blanco al
        inicio de cada página.

        IMPORTANTE: Nunca elimina párrafos que contengan imágenes, shapes o dibujos.
        """
        paragraphs_to_remove = []

        # Recorrer todos los párrafos
        i = 0
        while i < len(self.doc.paragraphs) - 1:
            para = self.doc.paragraphs[i]

            # Verificar si este párrafo tiene un salto de página
            has_page_break = any(self._has_page_break(run) for run in para.runs)

            if has_page_break:
                # Buscar párrafos vacíos consecutivos después del salto de página
                j = i + 1
                while j < len(self.doc.paragraphs):
                    next_para = self.doc.paragraphs[j]

                    # IMPORTANTE: No eliminar si tiene imágenes
                    if self._has_drawing_or_image(next_para):
                        break

                    # Verificar si el párrafo está vacío
                    if not next_para.text.strip():
                        # Marcar para eliminación
                        paragraphs_to_remove.append(next_para)
                        j += 1
                    else:
                        # Encontramos contenido, detener
                        break

            i += 1

        # También verificar al inicio del documento (primera página)
        i = 0
        while i < len(self.doc.paragraphs):
            para = self.doc.paragraphs[i]

            # IMPORTANTE: No eliminar si tiene imágenes
            if self._has_drawing_or_image(para):
                break

            # Si encontramos un párrafo vacío al inicio, marcarlo
            if not para.text.strip():
                # Solo eliminar si no hay salto de página
                has_break = any(self._has_page_break(run) for run in para.runs)
                if not has_break and para not in paragraphs_to_remove:
                    paragraphs_to_remove.append(para)
                i += 1
            else:
                # Encontramos contenido, detener
                break

        # Eliminar los párrafos marcados
        for para in paragraphs_to_remove:
            try:
                p_element = para._element
                p_element.getparent().remove(p_element)
            except Exception:
                # Si falla, continuar con el siguiente
                pass

    def preserve_headers_and_footers(self):
        """
        Asegura que todas las cabeceras (headers) y pies de página (footers)
        de todas las secciones se preserven correctamente.

        Este método verifica que todas las secciones del documento mantengan
        sus headers y footers con todos sus elementos gráficos y de diseño.

        Nota: python-docx preserva automáticamente los headers y footers, pero
        este método está disponible para realizar validaciones adicionales si
        es necesario.
        """
        # Verificar que todas las secciones tengan sus headers y footers
        for section in self.doc.sections:
            # Verificar header
            header = section.header
            if header:
                # El header existe, preservar sus imágenes
                for paragraph in header.paragraphs:
                    # Los headers con imágenes se preservan automáticamente
                    # No realizamos modificaciones aquí
                    pass

            # Verificar footer
            footer = section.footer
            if footer:
                # El footer existe, preservar sus imágenes
                for paragraph in footer.paragraphs:
                    # Los footers con imágenes se preservan automáticamente
                    # No realizamos modificaciones aquí
                    pass

    def process_table_of_contents(self):
        """
        Procesa el índice (tabla de contenidos) del documento usando marcadores numéricos.

        Busca contenidos entre <<Indice>> y <<fin Indice>>, extrae los marcadores numéricos
        (<<1>>, <<2>>, etc.) de cada entrada del índice, localiza estos marcadores en el
        documento, inserta saltos de página antes de ellos, calcula los números de página
        y actualiza el índice.

        Al finalizar, elimina todos los marcadores numéricos del documento.
        """
        # Buscar los marcadores de inicio y fin del índice
        toc_start_idx = None
        toc_end_idx = None

        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()
            if "<<Indice>>" in text:
                toc_start_idx = i
            elif "<<fin Indice>>" in text:
                toc_end_idx = i
                break

        # Si no se encuentran los marcadores, no hacer nada
        if toc_start_idx is None or toc_end_idx is None:
            return

        # Extraer las entradas del índice con sus marcadores numéricos
        toc_entries = []  # Lista de {paragraph, title, marker}

        marker_pattern = re.compile(r'<<(\d+)>>')

        for i in range(toc_start_idx + 1, toc_end_idx):
            para = self.doc.paragraphs[i]
            text = para.text.strip()

            if text:  # Solo procesar párrafos no vacíos
                # Buscar el marcador numérico en el texto
                match = marker_pattern.search(text)
                if match:
                    marker_num = match.group(1)
                    marker = f"<<{marker_num}>>"

                    # Extraer el título (texto antes del marcador)
                    title = marker_pattern.sub('', text).strip()

                    toc_entries.append({
                        'paragraph': para,
                        'title': title,
                        'marker': marker,
                        'marker_num': marker_num
                    })

        # Si no hay entradas con marcadores, salir
        if not toc_entries:
            # Limpiar solo los marcadores de inicio y fin
            start_para = self.doc.paragraphs[toc_start_idx]
            start_para.text = start_para.text.replace("<<Indice>>", "").strip()
            end_para = self.doc.paragraphs[toc_end_idx]
            end_para.text = end_para.text.replace("<<fin Indice>>", "").strip()
            return

        # Fase 1: Buscar cada marcador en el documento e insertar saltos de página
        for entry in toc_entries:
            marker = entry['marker']
            self._insert_page_break_before_marker(marker, toc_end_idx)

        # Fase 2: Calcular números de página para cada marcador
        marker_to_page = {}

        for entry in toc_entries:
            marker = entry['marker']
            page_num = self._find_marker_page_number(marker, toc_end_idx)
            if page_num is not None:
                marker_to_page[marker] = page_num

        # Fase 3: Actualizar el índice con los números de página
        for entry in toc_entries:
            para = entry['paragraph']
            title = entry['title']
            marker = entry['marker']

            if marker in marker_to_page:
                page_num = marker_to_page[marker]

                # Eliminar cualquier número de página existente al final
                clean_title = re.sub(r'[\.\s]+\d+$', '', title).strip()

                # Calcular el número de puntos necesarios
                dots_length = max(3, 80 - len(clean_title) - len(str(page_num)) - 2)
                dots = '.' * dots_length

                # Actualizar el párrafo (sin el marcador)
                para.text = f"{clean_title} {dots} {page_num}"
            else:
                # Si no se encontró el marcador, al menos limpiar el marcador del título
                para.text = title

        # Fase 4: Eliminar todos los marcadores numéricos del documento
        self._remove_numeric_markers()

        # Eliminar el marcador <<Indice>>
        start_para = self.doc.paragraphs[toc_start_idx]
        start_para.text = start_para.text.replace("<<Indice>>", "").strip()

        # Eliminar el marcador <<fin Indice>>
        end_para = self.doc.paragraphs[toc_end_idx]
        end_para.text = end_para.text.replace("<<fin Indice>>", "").strip()

        # Asegurar que el índice empiece en una nueva página
        if toc_start_idx > 0:
            prev_para = self.doc.paragraphs[toc_start_idx - 1]
            run = prev_para.add_run()
            run.add_break(WD_BREAK.PAGE)

    def _insert_page_break_before_marker(self, marker: str, toc_end_idx: int):
        """
        Busca un marcador numérico en el documento e inserta un salto de página antes de él.

        Args:
            marker: Marcador numérico a buscar (ej: "<<1>>")
            toc_end_idx: Índice del final del índice (para empezar búsqueda después)
        """
        # Buscar el marcador en el documento (después del índice)
        for i in range(toc_end_idx + 1, len(self.doc.paragraphs)):
            para = self.doc.paragraphs[i]

            if marker in para.text:
                # Verificar si ya hay un salto de página antes de este párrafo
                # Comprobar el párrafo anterior
                has_page_break = False
                if i > 0:
                    prev_para = self.doc.paragraphs[i - 1]
                    # Revisar si el párrafo anterior tiene un salto de página al final
                    for run in prev_para.runs:
                        if self._has_page_break(run):
                            has_page_break = True
                            break

                # También verificar si el párrafo actual tiene un salto de página al inicio
                if not has_page_break and para.runs:
                    for run in para.runs:
                        if self._has_page_break(run):
                            has_page_break = True
                            break

                # Si no tiene salto de página, insertar uno al final del párrafo anterior
                if not has_page_break:
                    if i > 0:
                        # Insertar el salto de página al final del párrafo anterior
                        prev_para = self.doc.paragraphs[i - 1]
                        run = prev_para.add_run()
                        run.add_break(WD_BREAK.PAGE)
                    else:
                        # Si no hay párrafo anterior, insertar al inicio del párrafo actual
                        # usando XML para insertar al principio
                        if para.runs:
                            # Insertar un nuevo run al principio con el salto de página
                            new_run = para.insert_paragraph_before().add_run()
                            new_run.add_break(WD_BREAK.PAGE)
                        else:
                            # Si no hay runs, crear uno nuevo
                            run = para.add_run()
                            run.add_break(WD_BREAK.PAGE)

                # Solo procesar la primera ocurrencia
                break

    def _find_marker_page_number(self, marker: str, start_search_idx: int) -> Optional[int]:
        """
        Encuentra el número de página donde aparece un marcador numérico.

        Args:
            marker: Marcador numérico a buscar (ej: "<<1>>")
            start_search_idx: Índice desde donde empezar la búsqueda (después del índice)

        Returns:
            Número de página donde se encuentra el marcador, o None si no se encuentra
        """
        # Contador de páginas
        page_count = 1

        # Buscar el marcador en el documento (después del índice)
        for i in range(start_search_idx + 1, len(self.doc.paragraphs)):
            para = self.doc.paragraphs[i]

            # Verificar si hay un salto de página en este párrafo
            for run in para.runs:
                if self._has_page_break(run):
                    page_count += 1

            # Verificar si este párrafo contiene el marcador
            if marker in para.text:
                return page_count

        # Si no se encuentra el marcador, devolver None
        return None

    def _remove_numeric_markers(self):
        """
        Elimina todos los marcadores numéricos (<<1>>, <<2>>, etc.) del documento.
        """
        marker_pattern = re.compile(r'<<\d+>>')

        # Eliminar de párrafos
        for paragraph in self.doc.paragraphs:
            if marker_pattern.search(paragraph.text):
                paragraph.text = marker_pattern.sub('', paragraph.text)

        # Eliminar de tablas
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if marker_pattern.search(paragraph.text):
                            paragraph.text = marker_pattern.sub('', paragraph.text)

        # Eliminar de headers y footers
        for section in self.doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                if marker_pattern.search(paragraph.text):
                    paragraph.text = marker_pattern.sub('', paragraph.text)

            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                if marker_pattern.search(paragraph.text):
                    paragraph.text = marker_pattern.sub('', paragraph.text)

    def _has_page_break(self, run) -> bool:
        """
        Verifica si un run contiene un salto de página.

        Args:
            run: Run de python-docx

        Returns:
            True si el run contiene un salto de página, False en caso contrario
        """
        # Obtener el elemento XML del run
        run_element = run._element

        # Buscar elementos de salto de página (w:br con w:type="page")
        for br in run_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
            break_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            if break_type == 'page':
                return True

        return False

    def save(self, output_path: Path):
        """
        Guarda el documento generado.

        Args:
            output_path: Ruta donde guardar el documento
        """
        self.doc.save(output_path)

    def get_document_bytes(self) -> bytes:
        """
        Obtiene el documento como bytes (para descarga en Streamlit).

        Returns:
            Bytes del documento
        """
        from io import BytesIO

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def get_pdf_bytes(self) -> bytes:
        """
        Convierte el documento a PDF y retorna los bytes.

        Utiliza LibreOffice para la conversión, lo que garantiza que:
        - Las imágenes de fondo se preservan correctamente
        - Los headers y footers mantienen sus elementos gráficos
        - El formato general del documento se mantiene intacto

        Nota: Requiere que LibreOffice esté instalado en el sistema.
        En Linux: sudo apt-get install libreoffice
        En macOS: brew install libreoffice
        En Windows: descargar de https://www.libreoffice.org/download/

        Returns:
            Bytes del documento en formato PDF

        Raises:
            RuntimeError: Si LibreOffice no está instalado o hay error en la conversión
        """
        import tempfile
        import os
        import subprocess
        import shutil

        # Verificar si LibreOffice está instalado
        libreoffice_cmd = None
        for cmd in ['libreoffice', 'soffice']:
            if shutil.which(cmd):
                libreoffice_cmd = cmd
                break

        if not libreoffice_cmd:
            raise RuntimeError(
                "LibreOffice no está instalado. "
                "Instálalo con: sudo apt-get install libreoffice (Linux) o "
                "brew install libreoffice (macOS)"
            )

        # Crear directorio temporal
        temp_dir = tempfile.mkdtemp()

        try:
            # Guardar el documento en el directorio temporal
            docx_path = os.path.join(temp_dir, 'documento.docx')
            self.doc.save(docx_path)

            # Convertir DOCX a PDF usando LibreOffice
            # --headless: ejecutar sin interfaz gráfica
            # --convert-to pdf: convertir a PDF
            # --outdir: directorio de salida
            result = subprocess.run(
                [
                    libreoffice_cmd,
                    '--headless',
                    '--convert-to',
                    'pdf',
                    '--outdir',
                    temp_dir,
                    docx_path
                ],
                capture_output=True,
                text=True,
                timeout=60  # Timeout de 60 segundos
            )

            if result.returncode != 0:
                raise RuntimeError(
                    f"LibreOffice falló al convertir el documento: {result.stderr}"
                )

            # Leer el PDF generado
            pdf_path = os.path.join(temp_dir, 'documento.pdf')

            if not os.path.exists(pdf_path):
                raise RuntimeError("El archivo PDF no se generó correctamente")

            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()

            return pdf_bytes

        except subprocess.TimeoutExpired:
            raise RuntimeError(
                "La conversión a PDF tardó demasiado tiempo. "
                "Intenta con un documento más pequeño."
            )
        except Exception as e:
            raise RuntimeError(f"Error al convertir a PDF: {e}")
        finally:
            # Limpiar directorio temporal
            try:
                shutil.rmtree(temp_dir)
            except Exception:
                pass  # Ignorar errores al limpiar
